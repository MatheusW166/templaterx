from src.templaterx import TemplaterX
from .constants import TEMPLATES_DIR
from tests.helpers import paths as p, docx, template, hash as hs
from docx.shared import Inches
import pytest


@pytest.fixture
def docx_main(tmp_path: p.Path):
    return p.Paths(
        template=TEMPLATES_DIR / "subdoc_tpl.docx",
        out=tmp_path / "out.docx",
    )


def test_subdoc_incremental_render_support(docx_main, tmp_path):
    tplx = TemplaterX(docx_main.template)

    sd = tplx.new_subdoc()

    p = sd.add_paragraph("This is a sub-document inserted into a bigger one")
    p = sd.add_paragraph("It has been ")
    p.add_run("dynamically").style = "dynamic"
    p.add_run(" generated with python by using ")
    p.add_run("python-docx").italic = True
    p.add_run(" library")

    sd.add_heading("Heading, level 1", level=1)
    sd.add_paragraph("This is an Intense quote", style="IntenseQuote")

    image_path = TEMPLATES_DIR / "python_logo.png"
    sd.add_paragraph("A picture :")
    sd.add_picture(str(image_path), width=Inches(1.25))

    sd.add_paragraph("A Table :")
    table = sd.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Qty"
    hdr_cells[1].text = "Id"
    hdr_cells[2].text = "Desc"
    recordset = (
        (1, 101, "Spam"),
        (2, 42, "Eggs"),
        (3, 631, "Spam,spam, eggs, and ham")
    )

    for item in recordset:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item[0])
        row_cells[1].text = str(item[1])
        row_cells[2].text = item[2]

    tplx.render({})
    tplx.render({"mysubdoc": sd})

    xml = docx.get_rendered_xml(tplx, docx_main.out)

    for text in [
        "This is a sub-document inserted into a bigger one",
        "It has been ",
        " generated with python by using ",
        " library",
        "This is an Intense quote",
        "A picture :",
        "A Table :",
    ]:
        assert f">{text}<" in xml

    template.assert_text_has_run_property(
        xml,
        "dynamically",
        prop_tag="w:rStyle",
        prop_attrs={"w:val": "dynamic"},
    )

    template.assert_text_has_run_property(
        xml,
        "python-docx",
        prop_tag="w:i",
        prop_attrs=None,
    )

    template.assert_text_has_paragraph_property(
        xml,
        "Heading, level 1",
        prop_tag="w:pStyle",
        prop_attrs={"w:val": "Titre1"},
    )

    template.assert_text_has_paragraph_property(
        xml,
        "This is an Intense quote",
        prop_tag="w:pStyle",
        prop_attrs={"w:val": "IntenseQuote"},
    )

    assert "<w:drawing>" in xml
    assert "<a:blip" in xml
    assert "image" in xml

    img = docx.extract_all_images(docx_main.out, tmp_path)[0]
    assert hs.file_hash(img) == hs.file_hash(image_path)

    assert "<w:tbl>" in xml
    assert "<w:tr>" in xml
    assert "<w:tc>" in xml

    for header in ["Qty", "Id", "Desc"]:
        assert f">{header}<" in xml

    for value in [
        "1", "101", "Spam",
        "2", "42", "Eggs",
        "3", "631", "Spam,spam, eggs, and ham",
    ]:
        assert f">{value}<" in xml


def test_subdoc_should_keep_placeholder(docx_main):

    tplx = TemplaterX(docx_main.template)
    tplx.render({})

    xml = docx.get_rendered_xml(tplx, docx_main.out)
    assert r"{{ mysubdoc }}" in xml
