from docx import Document
from pathlib import Path
from src.templaterx import TemplaterX
from lxml import etree as ET  # type: ignore


def extract_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)


def save_and_get_text(tplx: TemplaterX, tmp_path: Path) -> str:
    tplx.save(tmp_path)
    return extract_text(str(tmp_path))


def save_and_get_xml(tplx: TemplaterX, tmp_path: Path) -> str:
    tplx.save(tmp_path)
    doc = Document(str(tmp_path))
    return ET.tostring(doc._element.body, encoding="unicode", pretty_print=False)
